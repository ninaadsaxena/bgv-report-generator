#!/usr/bin/env python3
"""
BGV Report Generator — Core Engine  (v3.1)
============================================
Architecture v3.1:
  - Universal Semantic Header & Tag Matcher:
      Maps any Excel header variations (e.g. Candidate Name, Applicant, Full Name, DOB,
      Birth Date, Location, Address, Ref No, Check ID, Court, Authority, Source,
      Closure Date, Completion Date, Status) to their canonical semantic concepts.
  - Dynamic Template Tag Extraction:
      Scans any DOCX template to discover all <Tag> placeholders and automatically
      resolves each tag from either direct column matches or canonical concept matches.
  - Robust XML Run-Merging:
      Handles Word's internal split runs (e.g., '<', 'Candidate', ' Name>')
      with precise lxml tree manipulation and regex-based token replacement.
  - Multi-Sheet & Auto-Detection:
      Detects check type from sheet name or row metadata.
  - Zero-Disk Footprint:
      Keeps all generated PDFs in memory (bytes).
"""

import sys, os, re, shutil, datetime, argparse, json
import threading, tempfile
from pathlib import Path
from copy import deepcopy

import openpyxl

try:
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

SCRIPT_DIR    = Path(__file__).parent.resolve()
TEMPLATES_DIR = SCRIPT_DIR / "templates"

# ---------------------------------------------------------------------------
# Semantic Canonical Ontology & Concept Aliases
# ---------------------------------------------------------------------------

CANONICAL_CONCEPTS = {
    "NAME": {
        "aliases": [
            "candidate name", "applicant name", "candidate", "applicant",
            "employee name", "employee", "subject name", "subject",
            "person name", "full name", "client name", "name",
            "cand name", "app name", "cand_name", "applicant_name",
            "name of candidate", "name of applicant", "name of subject",
            "individual name", "profile name"
        ],
        "default": "Unknown",
        "is_name": True,
    },
    "DOB": {
        "aliases": [
            "date of birth", "dob", "d.o.b.", "birth date", "birthdate",
            "date_of_birth", "bday", "birthday", "d o b", "born date",
            "birth_date", "applicant dob", "candidate dob", "dob (dd/mm/yyyy)",
            "dob (yyyy-mm-dd)"
        ],
        "default": "N/A",
        "is_date": True,
    },
    "ADDRESS": {
        "aliases": [
            "complete address with country & zip/postal code (as applicable)",
            "complete address with country", "complete address", "full address",
            "candidate address", "applicant address", "current address",
            "permanent address", "residential address", "residence address",
            "address", "street address", "address line", "location",
            "residence", "place of residence", "verified address"
        ],
        "default": "N/A",
        "is_address": True,
    },
    "COUNTRY": {
        "aliases": [
            "country", "nation", "country name", "nationality",
            "country of residence", "country code"
        ],
        "default": "",
    },
    "CHECK_ID": {
        "aliases": [
            "check id", "check_id", "checkid", "case reference number",
            "case reference", "case ref number", "case ref no", "case ref",
            "caseref", "reference number", "reference no", "reference id",
            "ref no", "ref id", "case id", "case number", "case no",
            "check number", "check no", "report id", "ticket id",
            "case_reference_number", "check_no", "order id", "bgv id",
            "request id", "file number", "file no", "tracking id"
        ],
        "default": "N/A",
        "is_int_str": True,
    },
    "SOURCE": {
        "aliases": [
            "source name", "source", "relevant court", "court", "court name",
            "source_name", "verification source", "issuing authority",
            "authority", "police station", "jurisdiction", "court/authority",
            "agency", "vendor", "institution", "organization", "data source",
            "verification agency", "verified by", "court of jurisdiction"
        ],
        "default": "N/A",
    },
    "CLOSURE_DATE": {
        "aliases": [
            "closure date", "close date", "completion date", "completed date",
            "report date", "date of report", "verification date", "verified date",
            "end date", "closed date", "finish date", "closure_date", "close_date",
            "report issuance date", "issuance date", "date of closure", "issue date"
        ],
        "default": "N/A",
        "is_date": True,
    },
    "RECEIVED_DATE": {
        "aliases": [
            "case recived date", "case received date", "received date",
            "request date", "initiation date", "start date", "created date",
            "entry date", "order date", "case initiation date", "intimation date"
        ],
        "default": "N/A",
        "is_date": True,
    },
    "STATUS": {
        "aliases": [
            "status", "verification status", "result", "outcome",
            "disposition", "conclusion", "finding", "findings", "verdict",
            "case status", "check status", "final status", "record status"
        ],
        "default": "No records found",
    },
    "SEVERITY": {
        "aliases": [
            "severity", "risk level", "color", "colour", "alert level",
            "severity level", "risk"
        ],
        "default": "Green",
    },
    "CHECK_NAME": {
        "aliases": [
            "check name", "check type", "verification type", "service name",
            "service", "component", "check_name", "check_type", "type of check"
        ],
        "default": "Verification",
    },
    "REMARKS": {
        "aliases": [
            "remarks", "comments", "notes", "summary", "details", "additional notes"
        ],
        "default": "Verification completed. No adverse records found.",
    },
    "ID_TYPE": {
        "aliases": [
            "id type", "identity type", "document type", "doc type", "id document",
            "card type", "type of id", "gov id type", "identification type"
        ],
        "default": "Government ID",
    },
    "ID_NUMBER": {
        "aliases": [
            "id number", "id no", "identity number", "doc number", "document number",
            "doc no", "passport no", "passport number", "ssn", "national id",
            "national id number", "aadhar", "aadhaar", "pan", "pan number",
            "license number", "dl number", "voter id", "id_number", "id_no"
        ],
        "default": "N/A",
        "is_int_str": True,
    },
    "GENDER": {
        "aliases": ["gender", "sex"],
        "default": "N/A",
    },
    "FATHER_NAME": {
        "aliases": ["father name", "father's name", "parent name", "guardian name"],
        "default": "N/A",
    },
    "ISSUE_DATE": {
        "aliases": ["issue date", "date of issue", "issued date", "id issue date"],
        "default": "N/A",
        "is_date": True,
    },
    "EXPIRY_DATE": {
        "aliases": ["expiry date", "date of expiry", "expiration date", "valid till", "valid upto"],
        "default": "N/A",
        "is_date": True,
    },
}


def normalize_identifier(s):
    """Normalize string for fuzzy/semantic comparison: alphanumeric only, lowercased."""
    if not s:
        return ""
    s = str(s).strip().strip("<>").strip()
    return re.sub(r"[^a-z0-9]", "", s.lower())


def match_concept(identifier):
    """
    Match an identifier (from Excel header or template placeholder) to a canonical concept.
    Returns concept key ('NAME', 'DOB', etc.) or None.
    """
    norm = normalize_identifier(identifier)
    if not norm:
        return None

    # 1. Exact normalized alias match
    for concept, meta in CANONICAL_CONCEPTS.items():
        for alias in meta["aliases"]:
            if normalize_identifier(alias) == norm:
                return concept

    # 2. Substring overlap match (prefer longest match)
    best_concept = None
    best_len = 0
    for concept, meta in CANONICAL_CONCEPTS.items():
        for alias in meta["aliases"]:
            alias_norm = normalize_identifier(alias)
            if len(alias_norm) >= 3 and (alias_norm in norm or norm in alias_norm):
                match_len = min(len(alias_norm), len(norm))
                if match_len > best_len:
                    best_len = match_len
                    best_concept = concept

    return best_concept


# ---------------------------------------------------------------------------
# Utility helpers
# ---------------------------------------------------------------------------

def fmt_date(value, fallback="N/A"):
    """Format any date-like value to DD-Mon-YYYY."""
    if value is None or str(value).strip() == "":
        return fallback
    if isinstance(value, (datetime.datetime, datetime.date)):
        return value.strftime("%d-%b-%Y")
    val_str = str(value).strip()
    for fmt in ("%m/%d/%y", "%m/%d/%Y", "%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%d-%b-%Y", "%Y/%m/%d"):
        try:
            return datetime.datetime.strptime(val_str, fmt).strftime("%d-%b-%Y")
        except ValueError:
            continue
    return val_str


def clean(value, fallback="N/A"):
    """Strip and return string value, or fallback."""
    if value is None:
        return fallback
    s = str(value).strip()
    return s if s and s.lower() not in ("none", "nan", "n/a", "null") else fallback


def safe_int_str(value):
    """Convert float-like integers (e.g. 2161503951.0) to clean int strings."""
    if value is None:
        return "N/A"
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    s = str(value).strip()
    return s if s and s.lower() not in ("none", "nan", "n/a") else "N/A"


def make_report_filename(candidate_name, check_type):
    """
    Produce a short human-readable filename stem.
    'Bryan James Caisip Bermudez', 'Address'  →  'Bryan_Bermudez_Address_Report'
    """
    parts = str(candidate_name or "Unknown").strip().split()
    if len(parts) >= 2:
        short = f"{parts[0]}_{parts[-1]}"
    elif parts:
        short = parts[0]
    else:
        short = "Unknown"

    check_clean = re.sub(r"[^A-Za-z0-9]+", "_", str(check_type or "Report")).strip("_")
    name = f"{short}_{check_clean}_Report"
    name = re.sub(r"[^A-Za-z0-9_\-]+", "_", name)
    return re.sub(r"_+", "_", name).strip("_")


# ---------------------------------------------------------------------------
# Smart Row Context & Tag Resolver
# ---------------------------------------------------------------------------

def build_row_context(row_dict):
    """
    Indexes an Excel row both by direct normalized column name and by canonical concept.
    Also builds smart composite values (e.g. Address + Country).
    """
    direct_map = {}
    concept_map = {}

    for raw_col, val in row_dict.items():
        if val is None:
            continue
        s_val = str(val).strip()
        if not s_val or s_val.lower() in ("none", "nan", "null"):
            continue

        n_col = normalize_identifier(raw_col)
        direct_map[n_col] = val

        c = match_concept(raw_col)
        if c and c not in concept_map:
            concept_map[c] = val

    # Intelligent Composite: Address + Country
    addr_val = concept_map.get("ADDRESS") or direct_map.get(normalize_identifier("Address")) or ""
    country_val = concept_map.get("COUNTRY") or direct_map.get(normalize_identifier("Country")) or ""
    addr_str = clean(addr_val, fallback="")
    country_str = clean(country_val, fallback="")

    if country_str and country_str.lower() != "n/a" and country_str.lower() not in addr_str.lower():
        full_addr = f"{addr_str}, {country_str}".strip(", ") if addr_str else country_str
    else:
        full_addr = addr_str or country_str or "N/A"

    concept_map["ADDRESS"] = full_addr

    return direct_map, concept_map


def resolve_tag_value(tag, direct_map, concept_map, warnings=None):
    """
    Given a template placeholder tag (e.g. 'Applicant Name', 'Check Id', 'Closure date'),
    determine the best matching value from direct column matches or semantic concept mappings.
    """
    n_tag = normalize_identifier(tag)

    # 1. Semantic Concept Match
    concept = match_concept(tag)
    if concept and concept in concept_map:
        val = concept_map[concept]
        meta = CANONICAL_CONCEPTS.get(concept, {})
        if meta.get("is_date"):
            return fmt_date(val)
        elif meta.get("is_int_str"):
            return safe_int_str(val)
        return clean(val, fallback=meta.get("default", "N/A"))

    # 2. Direct Column Name Match
    if n_tag in direct_map:
        val = direct_map[n_tag]
        if isinstance(val, (datetime.datetime, datetime.date)):
            return fmt_date(val)
        return clean(val)

    # 3. Fallback Concept Default
    if concept:
        default_val = CANONICAL_CONCEPTS[concept].get("default", "N/A")
        return default_val

    if warnings is not None:
        warnings.append(f"Template tag <{tag}> could not be matched to any column.")
    return "N/A"


# ---------------------------------------------------------------------------
# python-docx placeholder substitution  (robust, handles split runs & regex)
# ---------------------------------------------------------------------------

def _replace_in_para(para, placeholder, replacement):
    """
    Replace <placeholder> with replacement text inside a paragraph.
    Handles split runs and flexible whitespace / case matching.
    """
    from lxml import etree
    W  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    NS = '{http://www.w3.org/XML/1998/namespace}'

    p_el = para._p
    runs = p_el.findall(f'{{{W}}}r')
    if not runs:
        return False

    # Build per-run text list
    run_texts = []
    for r in runs:
        t_els = r.findall(f'{{{W}}}t')
        run_texts.append("".join(t_el.text or "" for t_el in t_els))

    full = "".join(run_texts)
    
    # Regex match for tag with flexible whitespace and case-insensitivity
    pat = re.compile(rf'<\s*{re.escape(placeholder.strip())}\s*>', re.IGNORECASE)
    match = pat.search(full)
    if not match:
        return False

    tok_start = match.start()
    tok_end   = match.end()

    # Map each run to its [start, end) character offset in full
    offsets = []
    pos = 0
    for t in run_texts:
        offsets.append((pos, pos + len(t)))
        pos += len(t)

    # Categorize runs
    before_runs  = []
    during_runs  = []
    after_runs   = []

    for i, (start, end) in enumerate(offsets):
        if end <= tok_start:
            before_runs.append(i)
        elif start >= tok_end:
            after_runs.append(i)
        else:
            during_runs.append(i)

    if not during_runs:
        return False

    first_during = during_runs[0]
    last_during  = during_runs[-1]

    during_full = full[offsets[first_during][0] : offsets[last_during][1]]
    within_start = tok_start - offsets[first_during][0]
    within_end   = within_start + (tok_end - tok_start)
    new_during_text = (
        during_full[:within_start]
        + str(replacement)
        + during_full[within_end:]
    )

    # Rewrite the first during-run
    first_r = runs[first_during]
    for old_t in first_r.findall(f'{{{W}}}t'):
        first_r.remove(old_t)
    t_new = etree.SubElement(first_r, f'{{{W}}}t')
    t_new.text = new_during_text
    if new_during_text and (new_during_text[0] == ' ' or new_during_text[-1] == ' '):
        t_new.set(f'{NS}space', 'preserve')

    # Remove subsequent during-runs
    for i in during_runs[1:]:
        p_el.remove(runs[i])

    return True


def _extract_doc_placeholders(doc):
    """Scan all paragraphs, tables, headers, and footers to extract all unique <placeholder> tags."""
    placeholders = set()
    pat = re.compile(r'<([^<>]+)>')

    def check_text(text):
        for m in pat.findall(text):
            cleaned = m.strip()
            if cleaned:
                placeholders.add(cleaned)

    for p in doc.paragraphs:
        check_text("".join(r.text for r in p.runs))
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    check_text("".join(r.text for r in p.runs))
    for s in doc.sections:
        for hf in (s.header, s.footer):
            if hf:
                for p in hf.paragraphs:
                    check_text("".join(r.text for r in p.runs))

    return placeholders


def fill_docx_template(template_path, fields_or_row, out_docx_path, warnings=None):
    """
    Open template .docx, extract all <Placeholder> tags, resolve their values
    from fields_or_row using direct column matches + semantic canonical matching,
    and write to out_docx_path.
    """
    if not HAS_DOCX:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    if warnings is None:
        warnings = []

    doc = Document(str(template_path))

    # Build context from row or dictionary
    direct_map, concept_map = build_row_context(fields_or_row)

    # Discover all placeholders actually present in the template
    doc_placeholders = _extract_doc_placeholders(doc)

    # Build replacement mapping for all detected placeholders
    replacements = {}
    for tag in doc_placeholders:
        val = resolve_tag_value(tag, direct_map, concept_map, warnings)
        replacements[tag] = val

    # Also include any explicit keys passed directly in fields_or_row
    for k, v in fields_or_row.items():
        if k not in replacements:
            replacements[k] = clean(v)

    def _replace_in_all(placeholder, value):
        val_str = str(value)
        found = False
        for para in doc.paragraphs:
            if _replace_in_para(para, placeholder, val_str):
                found = True
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if _replace_in_para(para, placeholder, val_str):
                            found = True
        for section in doc.sections:
            for hdr in (section.header, section.footer):
                if hdr:
                    for para in hdr.paragraphs:
                        if _replace_in_para(para, placeholder, val_str):
                            found = True
        return found

    for tag, val in replacements.items():
        _replace_in_all(tag, val)

    doc.save(str(out_docx_path))


# ---------------------------------------------------------------------------
# PDF conversion — Word COM (primary) + LibreOffice CLI (fallback)
# ---------------------------------------------------------------------------

_word_lock = threading.Lock()


def convert_to_pdf_word_com(docx_path, out_dir):
    """Use Microsoft Word's ExportAsFixedFormat (Windows only, best quality)."""
    import win32com.client, pythoncom

    docx_path = Path(docx_path).resolve()
    out_dir   = Path(out_dir).resolve()
    pdf_path  = out_dir / (docx_path.stem + ".pdf")

    pythoncom.CoInitialize()
    word = None
    try:
        with _word_lock:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            doc = word.Documents.Open(str(docx_path))
            try:
                doc.ExportAsFixedFormat(
                    OutputFileName=str(pdf_path),
                    ExportFormat=17,
                    OpenAfterExport=False,
                    OptimizeFor=0,
                    Range=0,
                    Item=0,
                    IncludeDocProps=True,
                    KeepIRM=True,
                    CreateBookmarks=0,
                    DocStructureTags=True,
                    BitmapMissingFonts=True,
                    UseISO19005_1=False,
                )
            finally:
                doc.Close(False)
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()

    if not pdf_path.exists():
        raise RuntimeError(f"Word COM produced no PDF for {docx_path.name}")
    return pdf_path


def convert_to_pdf_libreoffice(docx_path, out_dir):
    """LibreOffice CLI fallback — works on Linux/macOS servers."""
    import subprocess
    docx_path = Path(docx_path).resolve()
    out_dir   = Path(out_dir).resolve()
    soffice_candidates = ["soffice", "libreoffice"]
    cmd = None
    for candidate in soffice_candidates:
        try:
            subprocess.run([candidate, "--version"], capture_output=True, timeout=10, check=True)
            cmd = candidate
            break
        except Exception:
            continue
    if not cmd:
        raise RuntimeError("Neither 'soffice' nor 'libreoffice' found on PATH.")
    result = subprocess.run(
        [cmd, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
        capture_output=True, text=True, timeout=120,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice failed: {result.stderr}")
    pdf_path = out_dir / (docx_path.stem + ".pdf")
    if not pdf_path.exists():
        raise RuntimeError("LibreOffice produced no PDF")
    return pdf_path


def convert_to_pdf(docx_path, out_dir):
    """Try win32com first; fall back to LibreOffice CLI."""
    try:
        return convert_to_pdf_word_com(docx_path, out_dir)
    except ImportError:
        pass
    except Exception as e:
        import logging
        logging.warning(f"win32com PDF conversion failed ({e}), trying LibreOffice…")
    return convert_to_pdf_libreoffice(docx_path, out_dir)


# ---------------------------------------------------------------------------
# Excel reading
# ---------------------------------------------------------------------------

def read_sheet(xlsx_path, sheet_name=None):
    """
    Read an Excel sheet and return (headers, data_rows).
    data_rows is a list of dicts keyed by header name.
    """
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb[sheet_name] if sheet_name else wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        raise ValueError(f"Sheet '{ws.title}' is empty.")
    headers = [str(h).strip() if h is not None else "" for h in rows[0]]
    data_rows = []
    for raw in rows[1:]:
        if raw is None or all(v is None for v in raw):
            continue
        data_rows.append(dict(zip(headers, raw)))
    return headers, data_rows


def read_all_sheets(xlsx_path):
    """Return {sheet_name: (headers, data_rows)} for all sheets."""
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    result = {}
    for name in wb.sheetnames:
        ws = wb[name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            result[name] = ([], [])
            continue
        headers = [str(h).strip() if h is not None else "" for h in rows[0]]
        data_rows = []
        for raw in rows[1:]:
            if raw is None or all(v is None for v in raw):
                continue
            data_rows.append(dict(zip(headers, raw)))
        result[name] = (headers, data_rows)
    return result


# ---------------------------------------------------------------------------
# Field builders (Backward Compatibility Wrappers)
# ---------------------------------------------------------------------------

def _get_candidate_name(row):
    """Helper to get candidate name from row."""
    d_map, c_map = build_row_context(row)
    return c_map.get("NAME") or "Unknown"


def build_address_fields(row, warnings=None):
    """Backward compatibility helper."""
    d_map, c_map = build_row_context(row)
    return {
        "Applicant Name":        c_map.get("NAME", "Unknown"),
        "Date of Birth":         fmt_date(c_map.get("DOB")),
        "Address":               c_map.get("ADDRESS", "N/A"),
        "Case Reference Number": safe_int_str(c_map.get("CHECK_ID")),
        "Relevant Court":        clean(c_map.get("SOURCE")),
        "Closure Date":          fmt_date(c_map.get("CLOSURE_DATE")),
    }


def build_criminal_fields(row, warnings=None):
    """Backward compatibility helper."""
    d_map, c_map = build_row_context(row)
    return {
        "Candidate Name": c_map.get("NAME", "Unknown"),
        "Date of Birth":  fmt_date(c_map.get("DOB")),
        "Address":        c_map.get("ADDRESS", "N/A"),
        "Check Id":       safe_int_str(c_map.get("CHECK_ID")),
        "Relevant Court": clean(c_map.get("SOURCE")),
        "Closure date":   fmt_date(c_map.get("CLOSURE_DATE")),
    }


def build_criminal_civil_fields(row, warnings=None):
    """Backward compatibility helper."""
    return build_criminal_fields(row, warnings)


# ---------------------------------------------------------------------------
# Sheet-name → check type detection
# ---------------------------------------------------------------------------

def detect_check_type(sheet_name):
    """
    Map sheet name to check type key.
    Returns one of: 'address', 'criminal', 'civil', 'criminal_civil', 'id', or None.
    """
    name = sheet_name.strip().lower()
    
    # 1. Criminal + Civil combined
    if any(x in name for x in ("crim & civil", "criminal & civil", "criminal+civil",
                                "criminal + civil", "crim+civil", "crim &amp; civil",
                                "crim & civ", "crim/civil", "criminal/civil", "crim civil")):
        return "criminal_civil"
    
    # 2. Civil standalone
    if any(x in name for x in ("civil", "litigation")) and not any(x in name for x in ("crim", "criminal")):
        return "civil"

    # 3. Address verification
    if any(x in name for x in ("address", "addr", "residence", "location")):
        return "address"

    # 4. Criminal record check
    if any(x in name for x in ("criminal", "crim", "crime", "police", "court")):
        return "criminal"

    # 5. ID / Identity verification
    if any(x in name for x in ("id", "identity", "national id", "passport", "aadhar", "aadhaar",
                                "ssn", "pan", "dl", "driving", "voter", "gov id", "identification")):
        return "id"

    return None


# ---------------------------------------------------------------------------
# Check-type registry & Smart Fuzzy Template Finder
# ---------------------------------------------------------------------------

def _find_template(check_type_key, candidate_names=None):
    """
    Intelligently find a template docx file:
    1. Direct candidate filename matches in templates/ or root.
    2. Fuzzy keyword search across all .docx files in templates/ and root.
    3. Fallback to any valid docx in templates/ if only one exists.
    """
    if candidate_names:
        for name in candidate_names:
            for parent in (TEMPLATES_DIR, SCRIPT_DIR):
                p = parent / name
                if p.exists() and not p.name.startswith("~$"):
                    return p

    # Scan available docx files
    all_docx = []
    for parent in (TEMPLATES_DIR, SCRIPT_DIR):
        if parent.exists():
            for f in parent.glob("*.docx"):
                if not f.name.startswith("~$"):
                    all_docx.append(f)

    # Keywords per check type
    kw_map = {
        "address": ["address", "addr", "residence"],
        "criminal_civil": ["civil", "crim"],
        "civil": ["civil", "litigation"],
        "criminal": ["criminal", "crim", "crime", "court", "police"],
        "id": ["id", "identity", "national", "passport", "gov"],
    }
    keywords = kw_map.get(check_type_key, [check_type_key])

    # 1. Combined criminal & civil
    if check_type_key == "criminal_civil":
        for doc in all_docx:
            low = doc.stem.lower()
            if ("crim" in low and "civil" in low) or "crim_civil" in low:
                return doc

    # 2. Standalone civil (must not be criminal+civil combined if civil standalone requested)
    if check_type_key == "civil":
        for doc in all_docx:
            low = doc.stem.lower()
            if "civil" in low and "crim" not in low:
                return doc

    # 3. Match any keyword
    for doc in all_docx:
        low = doc.stem.lower()
        if any(kw in low for kw in keywords):
            return doc

    # 4. Fallback: if there's any docx file in TEMPLATES_DIR, return the first
    templates_docx = [f for f in TEMPLATES_DIR.glob("*.docx") if not f.name.startswith("~$")]
    if templates_docx:
        return templates_docx[0]

    return None


CHECK_TYPE_REGISTRY = {
    "address": {
        "get_template":    lambda: _find_template("address", [
                               "Address_report_format.docx",
                               "Address report format.docx",
                               "Address_Report_Format.docx",
                               "Address.docx",
                               "Address_template.docx",
                           ]),
        "build_fields":    build_address_fields,
        "display_name":    "Address Verification",
    },
    "criminal": {
        "get_template":    lambda: _find_template("criminal", [
                               "Criminal_report_format.docx",
                               "Criminal report format.docx",
                               "Criminal.docx",
                               "Criminal and Civil report format.docx",
                               "Criminal_and_Civil_report_format.docx",
                           ]),
        "build_fields":    build_criminal_fields,
        "display_name":    "Criminal Record Check",
    },
    "civil": {
        "get_template":    lambda: _find_template("civil", [
                               "Civil_report_format.docx",
                               "Civil report format.docx",
                               "Civil.docx",
                               "Civil_Litigation_format.docx",
                               "Criminal and Civil report format.docx",
                           ]),
        "build_fields":    build_criminal_fields,
        "display_name":    "Civil Litigation Check",
    },
    "criminal_civil": {
        "get_template":    lambda: _find_template("criminal_civil", [
                               "Criminal and Civil report format.docx",
                               "Criminal_and_Civil_report_format.docx",
                               "Criminal_Civil_report_format.docx",
                               "Criminal_Civil.docx",
                               "Criminal_report_format.docx",
                           ]),
        "build_fields":    build_criminal_civil_fields,
        "display_name":    "Criminal + Civil Check",
    },
    "id": {
        "get_template":    lambda: _find_template("id", [
                               "ID_report_format.docx",
                               "ID report format.docx",
                               "Identity_report_format.docx",
                               "ID_Verification_format.docx",
                               "ID.docx",
                               "National_ID_format.docx",
                           ]),
        "build_fields":    build_address_fields,
        "display_name":    "Identity Verification",
    },
}


# ---------------------------------------------------------------------------
# Single-row generation
# ---------------------------------------------------------------------------

def generate_one(row, check_type_key, template_override, work_dir, warnings, row_idx):
    """
    Generate a single PDF from one data row.
    Returns (pdf_bytes, candidate_name, filename, display_name).
    """
    config = CHECK_TYPE_REGISTRY.get(check_type_key, CHECK_TYPE_REGISTRY["address"])

    template_path = Path(template_override) if template_override else config["get_template"]()
    if template_path is None or not template_path.exists():
        raise FileNotFoundError(
            f"Template for check type '{check_type_key}' not found. "
            "Upload a .docx template via the Templates tab."
        )

    # Candidate Name resolution
    d_map, c_map = build_row_context(row)
    candidate = c_map.get("NAME") or "Unknown"
    display_name = config["display_name"]
    base_name = make_report_filename(candidate, display_name)
    pdf_name  = f"{base_name}.pdf"

    docx_out  = work_dir / f"{base_name}_{row_idx}.docx"

    # Universal Semantic Tag Resolution & DOCX filling
    fill_docx_template(template_path, row, docx_out, warnings)

    pdf_path  = convert_to_pdf(docx_out, work_dir)
    pdf_bytes = pdf_path.read_bytes()

    # Clean up temp DOCX and PDF immediately
    pdf_path.unlink(missing_ok=True)
    docx_out.unlink(missing_ok=True)

    return pdf_bytes, candidate, pdf_name, display_name


# ---------------------------------------------------------------------------
# Main generation pipeline — processes ALL sheets, or one sheet
# ---------------------------------------------------------------------------

def generate_all(xlsx_path, sheet_name=None, template_override=None,
                 progress_callback=None):
    """
    Generate one PDF per data row across all sheets (or a single named sheet).

    Returns:
      {
        "generated": [
          {"filename": str, "pdf_bytes": bytes, "applicant": str,
           "check_name": str, "sheet": str, "row": int, "warnings": list}
        ],
        "errors": [str]
      }

    PDF bytes are held in memory; no output directory is written.
    """
    work_dir = Path(tempfile.mkdtemp(prefix="bgv_work_"))
    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
        sheet_names = [sheet_name] if sheet_name else wb.sheetnames

        all_sheet_data = {}
        total_rows = 0
        for sn in sheet_names:
            try:
                _, rows = read_sheet(xlsx_path, sn)
                all_sheet_data[sn] = rows
                total_rows += len(rows)
            except Exception as e:
                all_sheet_data[sn] = []

        report = {"generated": [], "errors": []}
        done_count = 0

        for sn in sheet_names:
            rows = all_sheet_data.get(sn, [])
            if not rows:
                report["errors"].append(f"Sheet '{sn}': no data rows found.")
                continue

            # Auto-detect check type from sheet name
            check_type_key = detect_check_type(sn)
            if check_type_key is None:
                for r in rows:
                    cn = str(r.get("Check Name", "") or "").strip().lower()
                    if cn:
                        check_type_key = detect_check_type(cn)
                        break

            if check_type_key is None:
                report["errors"].append(
                    f"Sheet '{sn}': cannot determine check type from sheet name. "
                    "Rename sheet to 'Address', 'Crim', or 'Crim & Civil'."
                )
                done_count += len(rows)
                if progress_callback:
                    progress_callback(done_count, total_rows, f"Skipped sheet '{sn}': unknown type")
                continue

            for i, row in enumerate(rows, start=1):
                row_warnings = []
                candidate = _get_candidate_name(row)
                s_no = row.get("S.No.", i)

                if progress_callback:
                    progress_callback(done_count, total_rows,
                                      f"Filling template for {candidate} ({sn})…")
                try:
                    pdf_bytes, cand_name, pdf_name, check_display = generate_one(
                        row, check_type_key, template_override, work_dir, row_warnings, i
                    )

                    if progress_callback:
                        progress_callback(done_count, total_rows,
                                          f"Converting to PDF for {candidate}…")

                    report["generated"].append({
                        "row":        i,
                        "s_no":       s_no,
                        "sheet":      sn,
                        "applicant":  cand_name,
                        "check_name": check_display,
                        "filename":   pdf_name,
                        "pdf_bytes":  pdf_bytes,
                        "warnings":   row_warnings,
                    })

                    done_count += 1
                    if progress_callback:
                        progress_callback(done_count, total_rows,
                                          f"Done: {cand_name}")

                except Exception as e:
                    done_count += 1
                    report["errors"].append(
                        f"Sheet '{sn}' row {i} (S.No. {s_no}, {candidate}): "
                        f"{type(e).__name__}: {e}"
                    )
                    if progress_callback:
                        progress_callback(done_count, total_rows,
                                          f"Error: {candidate} — {e}")

    finally:
        shutil.rmtree(work_dir, ignore_errors=True)

    return report


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Generate BGV PDF reports from an Excel file.")
    parser.add_argument("xlsx_path",  help="Path to the input .xlsx file")
    parser.add_argument("output_dir", help="Directory to write generated PDFs into")
    parser.add_argument("--sheet",    default=None, help="Sheet name (default: all sheets)")
    parser.add_argument("--template", default=None, help="Override template .docx path")
    args = parser.parse_args()

    def cb(done, total, msg):
        print(f"  [{done}/{total}] {msg}")

    report = generate_all(
        args.xlsx_path,
        sheet_name=args.sheet,
        template_override=args.template,
        progress_callback=cb,
    )

    out = Path(args.output_dir)
    out.mkdir(parents=True, exist_ok=True)
    for g in report["generated"]:
        (out / g["filename"]).write_bytes(g["pdf_bytes"])

    print(f"\nGenerated: {len(report['generated'])}   Errors: {len(report['errors'])}")
    for g in report["generated"]:
        flag = " ⚠" if g["warnings"] else ""
        print(f"  OK   [{g['sheet']}] row {g['row']:>3}  {g['applicant']:<35}  →  {g['filename']}{flag}")
        for w in g["warnings"]:
            print(f"         Warning: {w}")
    for e in report["errors"]:
        print(f"  FAIL {e}")
    if report["errors"]:
        sys.exit(1)


if __name__ == "__main__":
    main()
